#!/usr/bin/env python3

import subprocess
import json
import requests
from tabulate import tabulate
from datetime import datetime
from docx import Document

# === CONFIGURATION ===
DOC_OUTPUT_FILE = f"{datetime.now().strftime('%B-%d-%Y').lower()}-vulnerabilities_table.doc"
TXT_OUTPUT_FILE = "vulnerabilities_table.txt"
MATTERMOST_BOT_TOKEN = "gss9tpu31tn8fbn9wq1bq5g89h"
MATTERMOST_CHANNEL_ID = "8i7xqbnofbykjf8ky391stdqty"
MATTERMOST_API_URL = "https://mattermost.edusc.us/api/v4/posts"

# Get list of Docker images
try:
    result = subprocess.run(
        ["docker", "images", "--format", "{{.Repository}}:{{.Tag}}"],
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True
    )
    if result.returncode != 0:
        print(f"❌ Failed to get Docker images: {result.stderr.strip()}")
        exit(1)

    images = [line.strip() for line in result.stdout.splitlines() if line.strip()]
except Exception as e:
    print(f"❌ Error retrieving Docker images: {str(e)}")
    exit(1)

if not images:
    print("⚠️ No Docker images found.")
    exit(0)

summary_data = []
all_rows = []

# Scan each image
for image in images:
    print(f"🔍 Scanning image: {image}")
    scan_result = subprocess.run(
        ["/usr/bin/trivy", "image", "--scanners", "vuln", "--severity", "CRITICAL,HIGH", "--format", "json", image],
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True
    )

    if scan_result.returncode != 0:
        print(f"⚠️ Error scanning {image}: {scan_result.stderr.strip()}")
        continue

    try:
        data = json.loads(scan_result.stdout)
    except json.JSONDecodeError:
        print(f"❌ Failed to parse JSON for image {image}")
        continue

    critical_count = 0
    high_count = 0

    for target in data.get("Results", []):
        for vuln in target.get("Vulnerabilities", []):
            vuln_id = vuln.get("VulnerabilityID", "N/A")
            severity = vuln.get("Severity", "N/A")
            pkg_name = vuln.get("PkgName", "N/A")
            description = vuln.get("Title", "N/A")

            if description and len(description) > 80:
                description = description[:77] + "..."

            emoji = "🔴" if severity == "CRITICAL" else "🟠"
            all_rows.append([image, vuln_id, f"{emoji} {severity}", pkg_name, description])

            if severity == "CRITICAL":
                critical_count += 1
            elif severity == "HIGH":
                high_count += 1

    # Remove registry for shorter display
    short_image = image.split("/")[-1]
    summary_data.append([short_image, critical_count, high_count])

# Write full vulnerability table to text file
if all_rows:
    headers = ["Image", "CVE ID", "Severity", "Package", "Description"]
    table = tabulate(all_rows, headers=headers, tablefmt="grid")

    with open(TXT_OUTPUT_FILE, "w") as f:
        f.write(table)

    print(f"\n✅ Saved detailed vulnerability table to {TXT_OUTPUT_FILE}")
else:
    print("✅ No CRITICAL or HIGH vulnerabilities found.")

# Send summary table to Mattermost using bot
if summary_data:
    mm_headers = ["Image", "🔴 CRITICAL", "🟠 HIGH"]
    mm_table = tabulate(summary_data, headers=mm_headers, tablefmt="github")

    payload = {
        "channel_id": MATTERMOST_CHANNEL_ID,
        "message": (
            "@all\n\n"
            "🔐 **Connect Deployment Container Scan Summary**\n\n"
            f"{mm_table}\n\n"
            "📎 **Download the full CVE list here:** https://git.edusuc.net/WEBFORX/-/packages under **Assets**"
        )
    }

    headers = {
        "Authorization": f"Bearer {MATTERMOST_BOT_TOKEN}",
        "Content-Type": "application/json"
    }

    try:
        resp = requests.post(MATTERMOST_API_URL, headers=headers, json=payload)
        if resp.status_code == 201:
            print("📤 Summary posted to Mattermost via bot.")
        else:
            print(f"❌ Failed to post to Mattermost: {resp.status_code} - {resp.text}")
    except Exception as e:
        print(f"❌ Error sending message to Mattermost: {str(e)}")

# Generate Word doc report
if all_rows:
    doc = Document()
    doc.add_heading("Docker Image Vulnerability Report", level=1)
    doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}")

    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Image'
    hdr_cells[1].text = 'CVE ID'
    hdr_cells[2].text = 'Severity'
    hdr_cells[3].text = 'Package'
    hdr_cells[4].text = 'Description'

    for row in all_rows:
        image, vuln_id, severity, pkg, description = row
        cells = table.add_row().cells
        cells[0].text = image
        cells[1].text = vuln_id
        cells[2].text = severity
        cells[3].text = pkg
        cells[4].text = description

    doc.add_paragraph(
        "\nDownload the full CVE list here: https://git.edusuc.net/WEBFORX/-/packages under Assets"
    )

    doc.save(DOC_OUTPUT_FILE)
    print(f"📄 Saved detailed Word report to {DOC_OUTPUT_FILE}")
else:
    print("✅ No CRITICAL or HIGH vulnerabilities found.")
